import streamlit as st
import folium
from folium.plugins import Draw, Fullscreen
from streamlit_folium import st_folium
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import math
import pandas as pd
import json
from shapely.geometry import Polygon, Point
from shapely.ops import transform
import pyproj
from functools import partial
import os

# Try to import simplekml, but provide fallback if not available
try:
    import simplekml
    SIMPLEKML_AVAILABLE = True
except ImportError:
    SIMPLEKML_AVAILABLE = False
    st.sidebar.warning("KML export disabled. Install with: `pip install simplekml`")

def decimal_to_dms_formatted(decimal, is_lat):
    # Fix direction logic
    if is_lat:
        direction = "N" if decimal >= 0 else "S"
    else:
        direction = "E" if decimal >= 0 else "W"
    
    abs_val = abs(decimal)
    
    # Calculate degrees, minutes, and seconds
    deg = int(abs_val)
    min_full = (abs_val - deg) * 60
    min_val = int(min_full)
    sec = (min_full - min_val) * 60
    
    # Handle the case where seconds might be 60 due to floating point precision
    if sec >= 59.99995:  # Account for floating point errors
        sec = 0.0
        min_val += 1
        if min_val >= 60:
            min_val = 0
            deg += 1
    
    # Format with 2 digits for degrees and minutes
    deg_str = f"{deg:02d}"  # 2-digit degrees
    min_str = f"{min_val:02d}"  # 2-digit minutes
    
    # FIXED: Format seconds properly - separate integer and decimal parts
    sec_int = int(sec)
    sec_frac = sec - sec_int
    
    # Format fractional part to 4 decimal places
    sec_frac_formatted = f"{sec_frac:.4f}"[2:]  # Get the decimal part only
    if len(sec_frac_formatted) < 4:
        sec_frac_formatted = sec_frac_formatted.ljust(4, '0')
    
    # Combine integer and fractional parts
    sec_formatted = f"{sec_int:02d}.{sec_frac_formatted}"
    
    return f"{deg_str}°{min_str}'{sec_formatted}\"{direction}"


def calculate_distance(lat1, lon1, lat2, lon2):
    """Calculate distance between two coordinates in meters using Haversine formula"""
    try:
        R = 6371000  # Earth radius in meters
        
        # Validate inputs
        if not all(-90 <= x <= 90 for x in [lat1, lat2]) or not all(-180 <= x <= 180 for x in [lon1, lon2]):
            raise ValueError("Invalid coordinate values")
            
        lat1, lon1, lat2, lon2 = map(math.radians, [lat1, lon1, lat2, lon2])
        dlat = lat2 - lat1
        dlon = lon2 - lon1
        
        a = math.sin(dlat/2)**2 + math.cos(lat1) * math.cos(lat2) * math.sin(dlon/2)**2
        c = 2 * math.atan2(math.sqrt(a), math.sqrt(1-a))
        
        return R * c
    except Exception as e:
        st.error(f"Distance calculation error: {e}")
        return 0

def calculate_polyline_length(coordinates):
    """Calculate total length of a polyline in meters"""
    total_distance = 0
    for i in range(len(coordinates) - 1):
        lat1, lon1 = coordinates[i]
        lat2, lon2 = coordinates[i + 1]
        total_distance += calculate_distance(lat1, lon1, lat2, lon2)
    return total_distance

def calculate_polygon_area(coordinates):
    """Calculate area of a polygon in hectares using proper geodesic calculation"""
    if len(coordinates) < 3:
        return 0
    
    try:
        # Create a polygon from coordinates
        polygon = Polygon(coordinates)
        
        # Define the projection: WGS84 to equal area projection
        wgs84 = pyproj.CRS('EPSG:4326')
        # Use a suitable equal area projection for your region
        # For India, you might use UTM zone 43N, but for general use, World Mercator
        equal_area = pyproj.CRS('EPSG:3857')  # Web Mercator
        
        # Create transformer
        transformer = pyproj.Transformer.from_crs(wgs84, equal_area, always_xy=True)
        
        # Project the polygon and calculate area
        projected_polygon = transform(transformer.transform, polygon)
        area_sq_m = projected_polygon.area
        
        # Convert to hectares
        area_hectares = area_sq_m / 10000
        return area_hectares
        
    except Exception as e:
        st.warning(f"Area calculation warning: {e}. Using approximate method.")
        # Fallback to improved shoelace formula
        return calculate_polygon_area_approximate(coordinates)

def calculate_polygon_area_approximate(coordinates):
    """Improved approximate area calculation using Haversine"""
    if len(coordinates) < 3:
        return 0
    
    # Use spherical Earth approximation
    R = 6371000  # Earth radius in meters
    area = 0
    n = len(coordinates)
    
    for i in range(n):
        j = (i + 1) % n
        lat1, lon1 = coordinates[i]
        lat2, lon2 = coordinates[j]
        
        # Convert to radians
        lat1_rad = math.radians(lat1)
        lat2_rad = math.radians(lat2)
        lon1_rad = math.radians(lon1)
        lon2_rad = math.radians(lon2)
        
        # Spherical excess formula
        area += (lon2_rad - lon1_rad) * (2 + math.sin(lat1_rad) + math.sin(lat2_rad))
    
    area = abs(area) * R * R / 2.0
    return area / 10000  # Convert to hectares
def generate_waypoints_from_polyline_vertices(coordinates):
    """
    Generate waypoints from polyline vertices (each click/vertex becomes a waypoint)
    This is simpler and uses exactly the points you clicked
    """
    waypoints = []
    for coord in coordinates:
        waypoints.append({'lat': coord[0], 'lon': coord[1]})
    return waypoints

def parse_kml(file):
    try:
        content = file.read().decode('utf-8')
        
        try:
            root = ET.fromstring(content)
        except ET.ParseError:
            parser = ET.XMLParser(recover=True)
            root = ET.fromstring(content, parser=parser)
        
        coords = []
        namespaces = [
            {'kml': 'http://www.opengis.net/kml/2.2'},
            {'kml': 'http://earth.google.com/kml/2.0'},
            {'kml': 'http://earth.google.com/kml/2.1'},
            {}
        ]
        
        for ns in namespaces:
            try:
                # Look for coordinates in different possible locations
                coord_elements = []
                if ns:
                    coord_elements.extend(root.findall('.//{http://www.opengis.net/kml/2.2}coordinates', ns))
                    coord_elements.extend(root.findall('.//{http://www.opengis.net/kml/2.2}Point//{http://www.opengis.net/kml/2.2}coordinates', ns))
                    coord_elements.extend(root.findall('.//{http://www.opengis.net/kml/2.2}LinearRing//{http://www.opengis.net/kml/2.2}coordinates', ns))
                else:
                    coord_elements.extend(root.findall('.//coordinates'))
                    coord_elements.extend(root.findall('.//Point//coordinates'))
                    coord_elements.extend(root.findall('.//LinearRing//coordinates'))
                
                for elem in coord_elements:
                    if elem.text:
                        coord_text = elem.text.strip()
                        for coord in coord_text.split():
                            parts = coord.split(',')
                            if len(parts) >= 2:
                                try:
                                    lon = float(parts[0])
                                    lat = float(parts[1])
                                    # Validate coordinate ranges
                                    if -90 <= lat <= 90 and -180 <= lon <= 180:
                                        coords.append([lat, lon])  # Always [lat, lon]
                                    else:
                                        st.warning(f"Skipping invalid coordinates: {lat}, {lon}")
                                except ValueError:
                                    continue
                
                if coords:
                    break
                    
            except Exception as e:
                continue
        
        # Remove duplicates while preserving order
        seen = set()
        unique_coords = []
        for coord in coords:
            coord_tuple = (round(coord[0], 6), round(coord[1], 6))
            if coord_tuple not in seen:
                seen.add(coord_tuple)
                unique_coords.append(coord)
        
        return unique_coords
        
    except Exception as e:
        st.error(f"KML Parsing Error: {str(e)}")
        return []
def validate_coordinates(coordinates):
    """Validate that all coordinates are in proper [lat, lon] format"""
    valid_coords = []
    for coord in coordinates:
        if len(coord) == 2:
            lat, lon = coord
            if -90 <= lat <= 90 and -180 <= lon <= 180:
                valid_coords.append(coord)
            else:
                st.warning(f"Invalid coordinates skipped: {coord}")
    return valid_coords

def create_kml_manual(waypoints, date, kml_filename):
    """Create KML content manually without simplekml library"""
    kml_template = f"""<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://www.opengis.net/kml/2.2">
  <Document>
    <name>Flight Plan {date}</name>
    <description>Drone Survey Waypoints</description>
    """
    
    # Add waypoints
    for i, wp in enumerate(waypoints):
        kml_template += f"""
    <Placemark>
      <name>WP{chr(65+i)}</name>
      <description>Waypoint {chr(65+i)}</description>
      <Point>
        <coordinates>{wp['lon']},{wp['lat']},0</coordinates>
      </Point>
    </Placemark>
        """
    
    # Add flight path if multiple waypoints
    if len(waypoints) > 1:
        kml_template += """
    <Placemark>
      <name>Flight Path</name>
      <LineString>
        <coordinates>
        """
        for wp in waypoints:
            kml_template += f"          {wp['lon']},{wp['lat']},0\n"
        # Close the loop if more than 2 points
        if len(waypoints) > 2:
            kml_template += f"          {waypoints[0]['lon']},{waypoints[0]['lat']},0\n"
        kml_template += """
        </coordinates>
      </LineString>
      <Style>
        <LineStyle>
          <color>ff00ff00</color>
          <width>3</width>
        </LineStyle>
      </Style>
    </Placemark>
        """
    
    kml_template += """
  </Document>
</kml>
    """
    return kml_template

def ensure_lat_lon_order(coord):
    """Ensure coordinates are always in [lat, lon] order"""
    if len(coord) == 2:
        # If it's clearly lat, lon (reasonable values)
        if -90 <= coord[0] <= 90 and -180 <= coord[1] <= 180:
            return coord
        # If it's probably lon, lat
        elif -180 <= coord[0] <= 180 and -90 <= coord[1] <= 90:
            return [coord[1], coord[0]]
    return coord



# Initialize session state
if 'waypoints' not in st.session_state:
    st.session_state.waypoints = []
if 'kml_coords' not in st.session_state:
    st.session_state.kml_coords = []
if 'processed_markers' not in st.session_state:
    st.session_state.processed_markers = set()
if 'saved_projects' not in st.session_state:
    st.session_state.saved_projects = {}
if 'polylines' not in st.session_state:
    st.session_state.polylines = []
if 'active_polyline' not in st.session_state:
    st.session_state.active_polyline = None
if 'kml_filename' not in st.session_state:
    st.session_state.kml_filename = None
# Initialize layer_visibility
if 'layer_visibility' not in st.session_state:
    st.session_state.layer_visibility = {
        'waypoints': True,
        'flight_path': True,
        'kml_area': True,
        'saved_polylines': True,
        'active_polyline': True
    }

# =============================================================================
# DEFAULT SETTINGS - CUSTOMIZE THESE AS NEEDED:
# =============================================================================
# To change default base map, modify the DEFAULT_BASE_MAP value below:
# Options: "OpenStreetMap", "Satellite", "Terrain", "CartoDB"
DEFAULT_BASE_MAP = "OpenStreetMap"

# To change default active layers, modify the DEFAULT_LAYERS list below:
# Options: ["OpenStreetMap", "Satellite", "Terrain", "CartoDB"]
DEFAULT_LAYERS = ["OpenStreetMap", "Satellite"]
# =============================================================================

# 360°SURVEY V1 Drone Specifications
SURVEY_SPECS = {
    'max_flight_time': 30,      # minutes (from brochure: 25-30 min)
    'cruise_speed': 8.0,        # m/s (estimated for mapping operations)
    'battery_safety_margin': 15, # % (recommended safety buffer)
    'wind_resistance': 7.0,     # m/s (25 kmph = 6.94 m/s)
    'max_altitude': 120,        # meters
    'range': 1500,              # meters (1.5 km)
    'weight': 2.0,              # kg
    'accuracy_xy': 0.05,        # meters (<5cm)
    'accuracy_z': 0.10,         # meters (<10cm)
    'category': 'Micro'         # as per brochure
}

if 'drone_specs' not in st.session_state:
    st.session_state.drone_specs = SURVEY_SPECS.copy()

st.set_page_config(page_title="360°SURVEY Drone Survey Log", layout="wide")
st.title("🚁 360°SURVEY Drone Survey Log Generator")

# Display 360°SURVEY branding
st.sidebar.markdown("---")
st.sidebar.markdown("### 🛩️ 360°SURVEY V1")
st.sidebar.markdown("**Professional Survey Solutions**")

# Drone Specifications Section
with st.sidebar.expander("Drone Specifications", expanded=True):
    st.markdown("**Key Specs:**")
    st.markdown(f"- **Max Flight Time:** {st.session_state.drone_specs['max_flight_time']} min")
    st.markdown(f"- **Wind Resistance:** {st.session_state.drone_specs['wind_resistance']} m/s (25 kmph)")
    st.markdown(f"- **Max Altitude:** {st.session_state.drone_specs['max_altitude']} m")
    st.markdown(f"- **Accuracy:** XY: {st.session_state.drone_specs['accuracy_xy']*100}cm | Z: {st.session_state.drone_specs['accuracy_z']*100}cm")
    st.markdown(f"- **Range:** {st.session_state.drone_specs['range']} m")
    st.markdown(f"- **Weight:** {st.session_state.drone_specs['weight']} kg")
    st.markdown(f"- **Category:** {st.session_state.drone_specs['category']}")
    
    st.markdown("**Photogrammetry Ready**")
    st.markdown("- 80% overlap recommended")
    st.markdown("- Zig-zag flight patterns")
    st.markdown("- Fully autonomous waypoint missions")

# Map Settings in Sidebar
with st.sidebar.expander("Map Settings", expanded=False):
    # Find the index of the default base map for the selectbox
    map_options = ["OpenStreetMap", "Satellite", "Terrain", "CartoDB"]
    default_map_index = map_options.index(DEFAULT_BASE_MAP)
    
    default_base_map = st.selectbox(
        "Default Base Map",
        map_options,
        index=default_map_index,
        help="Set the default base map that loads when you open the app"
    )
    
    default_layers = st.multiselect(
        "Active Map Layers",
        ["OpenStreetMap", "Satellite", "Terrain", "CartoDB"],
        default=DEFAULT_LAYERS,
        help="Select which map layers should be available in the layer control"
    )

# TOP PANEL: Details & KML Upload
cols = st.columns([2,2,2,1,1,1])
with cols[0]:
    date = st.date_input("Date")
with cols[1]:
    pilot = st.text_input("Pilot", "Pretesh Ostwal")
with cols[2]:
    location = st.text_input("Location", "100 M BUFFER ZONE")
with cols[3]:
    altitude = st.number_input("Altitude (m)", value=80, min_value=1, max_value=120)
with cols[4]:
    start_time = st.text_input("Start Time", "1130", max_chars=4)
with cols[5]:
    end_time = st.text_input("End Time", "1146", max_chars=4)

st.divider()
kml = st.file_uploader("📍 Upload KML file", type=['kml'])
if kml:
    st.session_state.kml_coords = parse_kml(kml)
    # Extract filename without extension for use in downloads
    kml_filename = os.path.splitext(kml.name)[0]
    st.session_state.kml_filename = kml_filename
    st.success(f"KML loaded: {len(st.session_state.kml_coords)} points from {kml_filename}")

st.divider()

# SIDEBAR: All features
with st.sidebar:
    st.header("📍 Waypoints")
    
    # Waypoint list with delete buttons
    if st.session_state.waypoints:
        for i, wp in enumerate(st.session_state.waypoints):
            col1, col2 = st.columns([3, 1])
            with col1:
                st.write(f"**{chr(65+i)}**: {decimal_to_dms_formatted(wp['lat'], True)}, {decimal_to_dms_formatted(wp['lon'], False)}")
            with col2:
                if st.button("🗑️", key=f"del_{i}"):
                    st.session_state.waypoints.pop(i)
                    st.rerun()
    
    # Enhanced Waypoint Management (Removed optimization tools for photogrammetry)
    if st.session_state.waypoints:
        st.subheader("📊 Route Statistics")
        total_distance = 0
        for i in range(len(st.session_state.waypoints)):
            if i < len(st.session_state.waypoints) - 1:
                wp1 = st.session_state.waypoints[i]
                wp2 = st.session_state.waypoints[i + 1]
                distance = calculate_distance(wp1['lat'], wp1['lon'], wp2['lat'], wp2['lon'])
                total_distance += distance
        
        # Close the loop if more than 2 points
        if len(st.session_state.waypoints) > 2:
            wp1 = st.session_state.waypoints[-1]
            wp2 = st.session_state.waypoints[0]
            total_distance += calculate_distance(wp1['lat'], wp1['lon'], wp2['lat'], wp2['lon'])
        
        # Calculate statistics based on drone specs
        cruise_speed = st.session_state.drone_specs['cruise_speed']  # m/s
        max_flight_time = st.session_state.drone_specs['max_flight_time']  # minutes
        battery_margin = st.session_state.drone_specs['battery_safety_margin']  # %
        
        flight_time_minutes = total_distance / cruise_speed / 60
        battery_used = (flight_time_minutes / max_flight_time) * 100
        battery_remaining = max(0, 100 - battery_used)
        safe_battery_remaining = battery_remaining - battery_margin
        
        st.metric("Total Distance", f"{total_distance:.0f} m")
        st.metric("Number of Waypoints", len(st.session_state.waypoints))
        st.metric("Estimated Flight Time", f"{flight_time_minutes:.1f} min")
        
        # Battery status with color coding
        if safe_battery_remaining > 20:
            st.metric("Battery Used", f"{battery_used:.0f}%", delta=f"{battery_remaining:.0f}% remaining")
        elif safe_battery_remaining > 0:
            st.metric("Battery Used", f"{battery_used:.0f}%", delta=f"{battery_remaining:.0f}% remaining", delta_color="off")
        else:
            st.metric("Battery Used", f"{battery_used:.0f}%", delta="Insufficient battery", delta_color="inverse")
        
        # Flight feasibility check
        if battery_used > 100:
            st.error("❌ Route exceeds maximum flight time!")
        elif battery_used > (100 - battery_margin):
            st.warning("⚠️ Route uses most of available battery")
        else:
            st.success("✅ Route feasible within battery limits")
        
        # Photogrammetry specific information
        st.subheader("📷 Photogrammetry Info")
        estimated_photos = max(10, len(st.session_state.waypoints) * 3)  # Rough estimate
        st.metric("Estimated Photos", f"~{estimated_photos}")
        st.info("💡 For 80% overlap, maintain consistent altitude and parallel flight lines")
    
    else:
        st.info("Draw waypoints on map to see flight statistics")
    
    # Polyline Tools Section
    st.header("🔄 Polyline Tools")
    
    if st.session_state.polylines:
        st.subheader("Saved Polylines")
        for i, polyline in enumerate(st.session_state.polylines):
            col1, col2 = st.columns([3, 1])
            with col1:
                length = calculate_polyline_length(polyline)
                st.write(f"**Polyline {i+1}**: {len(polyline)} points, {length:.0f}m")
            with col2:
                if st.button("🗑️", key=f"poly_del_{i}"):
                    st.session_state.polylines.pop(i)
                    st.rerun()
    
    # Simplified Polyline generation - each vertex becomes a waypoint
    st.subheader("Generate Waypoints from Polyline")
    st.info("🎯 Each vertex/corner point in your polyline will become a waypoint")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🎯 Generate Waypoints", use_container_width=True, key="generate_waypoints"):
            if st.session_state.active_polyline and len(st.session_state.active_polyline) >= 2:
                new_waypoints = generate_waypoints_from_polyline_vertices(st.session_state.active_polyline)
                if new_waypoints:
                    st.session_state.waypoints = new_waypoints
                    st.session_state.processed_markers = set()
                    st.success(f"Generated {len(new_waypoints)} waypoints from polyline vertices!")
                    st.rerun()
            else:
                st.warning("No active polyline found. Draw a polyline first.")
    
    with col2:
        if st.button("💾 Save Polyline", use_container_width=True, key="save_polyline"):
            if st.session_state.active_polyline and len(st.session_state.active_polyline) >= 2:
                st.session_state.polylines.append(st.session_state.active_polyline)
                st.session_state.active_polyline = None
                st.success("Polyline saved!")
                st.rerun()
            else:
                st.warning("No active polyline to save.")
    
    if st.session_state.active_polyline:
        length = calculate_polyline_length(st.session_state.active_polyline)
        st.info(f"Active polyline: {len(st.session_state.active_polyline)} vertices, {length:.0f}m")
        st.write("**Note:** Each vertex will become a waypoint when you click 'Generate Waypoints'")
    
    st.divider()
    
    # Clear All Button
    if st.session_state.waypoints and st.button("🗑️ Clear All Waypoints", key="clear_waypoints"):
        st.session_state.waypoints = []
        st.session_state.processed_markers = set()
        st.rerun()
    
    st.divider()
    
    # Project Management
    st.header("💾 Project Management")
    project_name = st.text_input("Project Name", value=f"Survey_{date}")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 Save Project", key="save_project"):
            project_data = {
                'waypoints': st.session_state.waypoints,
                'date': str(date),
                'pilot': pilot,
                'location': location,
                'altitude': altitude,
                'start_time': start_time,
                'end_time': end_time,
                'kml_coords': st.session_state.kml_coords,
                'drone_specs': st.session_state.drone_specs,
                'polylines': st.session_state.polylines,
                'kml_filename': st.session_state.kml_filename
            }
            st.session_state.saved_projects[project_name] = project_data
            st.success(f"Project '{project_name}' saved!")
    
    with col2:
        if st.session_state.saved_projects:
            selected_project = st.selectbox("Load Project", list(st.session_state.saved_projects.keys()))
            if st.button("📂 Load Project", key="load_project"):
                project = st.session_state.saved_projects[selected_project]
                st.session_state.waypoints = project['waypoints']
                st.session_state.kml_coords = project.get('kml_coords', [])
                st.session_state.drone_specs = project.get('drone_specs', st.session_state.drone_specs)
                st.session_state.polylines = project.get('polylines', [])
                st.session_state.kml_filename = project.get('kml_filename', None)
                st.rerun()
    
    st.divider()
    
    # Export in Multiple Formats
    st.header("📤 Export Options")
    
    if len(st.session_state.waypoints) >= 2:
        # Word Document Export
        legs = []
        for i in range(len(st.session_state.waypoints)):
            from_wp = st.session_state.waypoints[i]
            to_wp = st.session_state.waypoints[(i + 1) % len(st.session_state.waypoints)]
            legs.append((from_wp, to_wp, i))
        
        if st.button("📝 Download Word (.docx)", use_container_width=True, key="download_word"):
            doc = Document()
            title = doc.add_paragraph('360°SURVEY DRONE SURVEY LOG SHEET')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.bold = True
            title.runs[0].font.size = Pt(14)
            doc.add_paragraph()
            section = doc.add_paragraph('2. Flight Log sheet')
            section.runs[0].font.bold = True
            section.runs[0].font.underline = True
            section.runs[0].font.size = Pt(11)
            doc.add_paragraph()
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            h_cells = table.rows[0].cells
            headers = ['S. No.', 'Date', 'Name of Remote Pilot', 'From', 'To', 'Place of Operation', 'Time of Operation', 'Height of Flight']
            for j, h in enumerate(headers):
                h_cells[j].text = h
                for p in h_cells[j].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for leg_idx, (from_wp, to_wp, i) in enumerate(legs):
                r_cells = table.add_row().cells
                if leg_idx == 0:
                    r_cells[0].text = "1"
                    r_cells[1].text = str(date)
                    r_cells[2].text = pilot
                    r_cells[5].text = location
                    r_cells[6].text = f"{start_time} HRS\n{end_time} HRS"
                    r_cells[7].text = f"{altitude}m AGL"
                else:
                    r_cells[0].text = ""
                    r_cells[1].text = ""
                    r_cells[2].text = ""
                    r_cells[5].text = ""
                    r_cells[6].text = ""
                    r_cells[7].text = ""
                from_lat = decimal_to_dms_formatted(from_wp['lat'], True)
                from_lon = decimal_to_dms_formatted(from_wp['lon'], False)
                r_cells[3].text = f"{from_lat}\n{from_lon}"
                to_lat = decimal_to_dms_formatted(to_wp['lat'], True)
                to_lon = decimal_to_dms_formatted(to_wp['lon'], False)
                r_cells[4].text = f"{to_lat}\n{to_lon}"
            
            # REMOVED: Drone specifications section
            # This section has been deleted as requested
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            # Use KML filename if available, otherwise use default
            if st.session_state.kml_filename:
                filename = f"{st.session_state.kml_filename}_{date}.docx"
            else:
                filename = f"{date}_360SURVEY_Log.docx"
                
            st.download_button(
                label="⬇️ Download .docx",
                data=output,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        
        # KML Export
        if SIMPLEKML_AVAILABLE:
            if st.button("🗺️ Export KML", use_container_width=True, key="export_kml"):
                kml = simplekml.Kml()
                
                # Add waypoints
                for i, wp in enumerate(st.session_state.waypoints):
                    pnt = kml.newpoint(name=f"WP{chr(65+i)}")
                    pnt.coords = [(wp['lon'], wp['lat'])]
                    pnt.style.iconstyle.scale = 1
                    pnt.style.iconstyle.color = simplekml.Color.red
                    pnt.description = f"Waypoint {chr(65+i)}\nLat: {wp['lat']:.6f}\nLon: {wp['lon']:.6f}"
                
                # Add flight path
                if len(st.session_state.waypoints) > 1:
                    linestring = kml.newlinestring(name="Flight Path")
                    coords = [(wp['lon'], wp['lat']) for wp in st.session_state.waypoints]
                    # Close the loop if more than 2 points
                    if len(st.session_state.waypoints) > 2:
                        coords.append(coords[0])
                    linestring.coords = coords
                    linestring.style.linestyle.color = simplekml.Color.green
                    linestring.style.linestyle.width = 3
                
                kml_data = kml.kml()
                
                # Use KML filename if available, otherwise use default
                if st.session_state.kml_filename:
                    filename = f"{st.session_state.kml_filename}_{date}.kml"
                else:
                    filename = f"{date}_360SURVEY_Flight_Plan.kml"
                    
                st.download_button(
                    label="⬇️ Download KML",
                    data=kml_data,
                    file_name=filename,
                    mime="application/vnd.google-earth.kml+xml",
                    use_container_width=True
                )
        else:
            # Manual KML export without simplekml
            if st.button("🗺️ Export KML (Basic)", use_container_width=True, key="export_kml_basic"):
                kml_data = create_kml_manual(st.session_state.waypoints, date, st.session_state.kml_filename)
                
                # Use KML filename if available, otherwise use default
                if st.session_state.kml_filename:
                    filename = f"{st.session_state.kml_filename}_{date}.kml"
                else:
                    filename = f"{date}_360SURVEY_Flight_Plan.kml"
                    
                st.download_button(
                    label="⬇️ Download KML (Basic)",
                    data=kml_data,
                    file_name=filename,
                    mime="application/vnd.google-earth.kml+xml",
                    use_container_width=True
                )
        
        # CSV Export
        if st.button("📊 Export CSV", use_container_width=True, key="export_csv"):
            waypoints_data = []
            for i, wp in enumerate(st.session_state.waypoints):
                waypoints_data.append({
                    'Waypoint': chr(65+i),
                    'Latitude': wp['lat'],
                    'Longitude': wp['lon'],
                    'Latitude_DMS': decimal_to_dms_formatted(wp['lat'], True),
                    'Longitude_DMS': decimal_to_dms_formatted(wp['lon'], False)
                })
            
            df = pd.DataFrame(waypoints_data)
            csv_data = df.to_csv(index=False)
            
            # Use KML filename if available, otherwise use default
            if st.session_state.kml_filename:
                filename = f"{st.session_state.kml_filename}_{date}.csv"
            else:
                filename = f"{date}_360SURVEY_Waypoints.csv"
                
            st.download_button(
                label="⬇️ Download CSV",
                data=csv_data,
                file_name=filename,
                mime="text/csv",
                use_container_width=True
            )
        
        # JSON Export (additional format)
        if st.button("📄 Export JSON", use_container_width=True, key="export_json"):
            export_data = {
                'project_info': {
                    'date': str(date),
                    'pilot': pilot,
                    'location': location,
                    'altitude': altitude,
                    'start_time': start_time,
                    'end_time': end_time
                },
                'drone_specs': st.session_state.drone_specs,
                'waypoints': st.session_state.waypoints,
                'statistics': {
                    'total_waypoints': len(st.session_state.waypoints),
                    'total_distance': total_distance,
                    'estimated_flight_time': flight_time_minutes,
                    'battery_usage': battery_used
                }
            }
            json_data = json.dumps(export_data, indent=2)
            
            # Use KML filename if available, otherwise use default
            if st.session_state.kml_filename:
                filename = f"{st.session_state.kml_filename}_{date}.json"
            else:
                filename = f"{date}_360SURVEY_Survey_Data.json"
                
            st.download_button(
                label="⬇️ Download JSON",
                data=json_data,
                file_name=filename,
                mime="application/json",
                use_container_width=True
            )
    else:
        st.warning("Add at least 2 waypoints to enable exports")

# MAIN PANEL
st.header("🗺️ Interactive Map")

# Real-time Coordinates Display
if st.session_state.waypoints:
    center = [st.session_state.waypoints[0]['lat'], st.session_state.waypoints[0]['lon']]
else:
    center = st.session_state.kml_coords[0] if st.session_state.kml_coords else [24.64, 72.58]

# Create two columns for map and layer panel with better ratio
map_col, layer_col = st.columns([3, 1])  # Changed from [4,1] to [3,1] for more map space

with map_col:
    # Real-time coordinates info
    st.info(f"**Map Center**: {center[0]:.6f}, {center[1]:.6f}")

    # Create the map with proper tile configuration
    m = folium.Map(location=center, zoom_start=16)

    # Define tile layers with proper attribution
    tile_layers = {
        "OpenStreetMap": folium.TileLayer(
            tiles='OpenStreetMap',
            attr='© OpenStreetMap contributors',
            name='OpenStreetMap',
            control=True
        ),
        "Satellite": folium.TileLayer(
            tiles='https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
            attr='Tiles &copy; Esri &mdash; Source: Esri, i-cubed, USDA, USGS, AEX, GeoEye, Getmapping, Aerogrid, IGN, IGP, UPR-EGP, and the GIS User Community',
            name='Satellite',
            control=True
        ),
        "Terrain": folium.TileLayer(
            tiles='Stamen Terrain',
            attr='Map tiles by Stamen Design, under CC BY 3.0. Data by OpenStreetMap, under ODbL.',
            name='Terrain',
            control=True
        ),
        "CartoDB": folium.TileLayer(
            tiles='CartoDB positron',
            attr='© OpenStreetMap contributors © CartoDB',
            name='CartoDB Light',
            control=True
        )
    }

    # Add the selected tile layer first (this becomes the default base map)
    tile_layers[default_base_map].add_to(m)

    # Add other tile layers based on user selection in settings
    for name, layer in tile_layers.items():
        if name != default_base_map and name in default_layers:
            layer.add_to(m)

    # Add layer control
    folium.LayerControl().add_to(m)

    # Add fullscreen button
    Fullscreen(
        position="topright",
        title="Full Screen",
        title_cancel="Exit Full Screen",
        force_separate_button=True,
    ).add_to(m)

    # Add KML polygon if available and visible
    if st.session_state.kml_coords and st.session_state.layer_visibility['kml_area']:
        folium.Polygon(
            locations=st.session_state.kml_coords,
            color='blue',
            fill=True,
            fillColor='lightblue',
            fillOpacity=0.2,
            weight=2.5,
            popup="Original Survey Area"
        ).add_to(m)

    # Add waypoints if visible
    if st.session_state.layer_visibility['waypoints']:
        for i, wp in enumerate(st.session_state.waypoints):
            folium.CircleMarker(
                location=[wp['lat'], wp['lon']],
                radius=8,
                color='red',
                fill=True,
                fillColor='red',
                fillOpacity=0.8,
                popup=f"Waypoint {chr(65+i)}",
                tooltip=f"{chr(65+i)}"
            ).add_to(m)

    # Add flight path if visible
    if st.session_state.layer_visibility['flight_path'] and len(st.session_state.waypoints) > 1:
        path = [[wp['lat'], wp['lon']] for wp in st.session_state.waypoints]
        if len(st.session_state.waypoints) > 2:
            path.append(path[0])  # Close the loop
        folium.PolyLine(path, color='green', weight=2.5, opacity=0.7, popup="Flight Path").add_to(m)

    # Add saved polylines if visible
    if st.session_state.layer_visibility['saved_polylines']:
        for i, polyline in enumerate(st.session_state.polylines):
            folium.PolyLine(
                locations=polyline,
                color='orange',
                weight=4,
                opacity=0.8,
                popup=f"Saved Polyline {i+1}",
                tooltip=f"Polyline {i+1}"
            ).add_to(m)

    # Add active polyline if visible
    if st.session_state.layer_visibility['active_polyline'] and st.session_state.active_polyline:
        folium.PolyLine(
            locations=st.session_state.active_polyline,
            color='purple',
            weight=6,
            opacity=0.9,
            popup="Active Polyline",
            tooltip="Active Polyline (Click Generate Waypoints to use)"
        ).add_to(m)

    # Advanced Drawing Tools
    Draw(
        export=False,
        position='topleft',
        draw_options={
            'polyline': True,
            'polygon': True,
            'rectangle': True,
            'circle': True,
            'marker': True,
            'circlemarker': False
        },
        edit_options={'edit': True, 'remove': True}
    ).add_to(m)

    # Display the map with larger size
    map_data = st_folium(m, width=1130, height=1100)  # Increased width and height

    # Display click coordinates
    if map_data and map_data.get('last_clicked'):
        lat = map_data['last_clicked']['lat']
        lon = map_data['last_clicked']['lng']
        st.info(f"**Last clicked**: {lat:.6f}, {lon:.6f} | **DMS**: {decimal_to_dms_formatted(lat, True)}, {decimal_to_dms_formatted(lon, False)}")

with layer_col:
    st.subheader("🗂️ Layer Panel")
    st.markdown("---")
    
    # Base Map Selection
    st.markdown("#### 🗺️ Base Maps")
    base_map = st.radio(
        "Select Base Map:",
        ["OpenStreetMap", "Satellite", "Terrain", "CartoDB"],
        index=["OpenStreetMap", "Satellite", "Terrain", "CartoDB"].index(default_base_map),
        key="base_map_selector"
    )
    
    st.markdown("---")
    
    # Layer Visibility Controls
    st.markdown("#### 👁️ Layer Visibility")
    
    # Waypoints layer
    waypoints_visible = st.checkbox(
        "📍 Waypoints", 
        value=st.session_state.layer_visibility['waypoints'],
        key="waypoints_visibility"
    )
    
    # Flight Path layer
    flight_path_visible = st.checkbox(
        "🛩️ Flight Path", 
        value=st.session_state.layer_visibility['flight_path'],
        key="flight_path_visibility"
    )
    
    # KML Area layer
    kml_visible = st.checkbox(
        "🔵 Survey Area (KML)", 
        value=st.session_state.layer_visibility['kml_area'],
        key="kml_visibility"
    )
    
    # Saved Polylines layer
    saved_polylines_visible = st.checkbox(
        "🟠 Saved Polylines", 
        value=st.session_state.layer_visibility['saved_polylines'],
        key="saved_polylines_visibility"
    )
    
    # Active Polyline layer
    active_polyline_visible = st.checkbox(
        "🟣 Active Polyline", 
        value=st.session_state.layer_visibility['active_polyline'],
        key="active_polyline_visibility"
    )
    
    # Update layer visibility in session state
    if (waypoints_visible != st.session_state.layer_visibility['waypoints'] or
        flight_path_visible != st.session_state.layer_visibility['flight_path'] or
        kml_visible != st.session_state.layer_visibility['kml_area'] or
        saved_polylines_visible != st.session_state.layer_visibility['saved_polylines'] or
        active_polyline_visible != st.session_state.layer_visibility['active_polyline']):
        
        st.session_state.layer_visibility['waypoints'] = waypoints_visible
        st.session_state.layer_visibility['flight_path'] = flight_path_visible
        st.session_state.layer_visibility['kml_area'] = kml_visible
        st.session_state.layer_visibility['saved_polylines'] = saved_polylines_visible
        st.session_state.layer_visibility['active_polyline'] = active_polyline_visible
        st.rerun()
    
    st.markdown("---")
    
    # Quick Actions - FIXED
    st.markdown("#### ⚡ Quick Actions")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("👁️ Show All", use_container_width=True, key="show_all_layers"):
            for key in st.session_state.layer_visibility:
                st.session_state.layer_visibility[key] = True
            st.rerun()
    
    with col2:
        if st.button("👁️ Hide All", use_container_width=True, key="hide_all_layers"):
            for key in st.session_state.layer_visibility:
                st.session_state.layer_visibility[key] = False
            st.rerun()
    
    st.markdown("---")
    
    # Layer Information
    st.markdown("#### ℹ️ Layer Info")
    
    if st.session_state.waypoints:
        st.write(f"**Waypoints:** {len(st.session_state.waypoints)}")
    
    if st.session_state.polylines:
        st.write(f"**Polylines:** {len(st.session_state.polylines)}")
    
    # Calculate and display KML area in hectares
    if st.session_state.kml_coords:
        if len(st.session_state.kml_coords) >= 3:
            area_hectares = calculate_polygon_area(st.session_state.kml_coords)
            st.write(f"**KML Area:** {area_hectares:.2f} HA.")
        else:
            st.write(f"**KML Points:** {len(st.session_state.kml_coords)}")
    
    if st.session_state.active_polyline:
        st.write(f"**Active Polyline:** {len(st.session_state.active_polyline)} vertices")

# Process drawings from map
if map_data and 'all_drawings' in map_data:
    drawings = map_data['all_drawings']
    if drawings:
        for drawing in drawings:
            # Process points (waypoints)
            if drawing['geometry']['type'] == 'Point':
                coords = drawing['geometry']['coordinates']
                lat, lon = coords[1], coords[0]
                
                # Create unique marker identifier
                marker_id = (round(lat, 6), round(lon, 6))
                
                # Only add if we haven't processed this marker before
                if marker_id not in st.session_state.processed_markers:
                    exists = any(abs(wp['lat'] - lat) < 0.0001 and abs(wp['lon'] - lon) < 0.0001 for wp in st.session_state.waypoints)
                    if not exists:
                        st.session_state.waypoints.append({'lat': lat, 'lon': lon})
                        st.session_state.processed_markers.add(marker_id)
                        st.rerun()
            
            # Process polylines
            elif drawing['geometry']['type'] == 'LineString':
                coords = drawing['geometry']['coordinates']
                # Convert from [lon, lat] to [lat, lon]
                polyline_coords = [[coord[1], coord[0]] for coord in coords]
                
                # Store as active polyline
                st.session_state.active_polyline = polyline_coords
                st.rerun()